import json

import docx
import pandas as pd

from preprocessing.pipeline import Phase2Pipeline
from preprocessing.validation import (
    validate_against_schema,
)


def test_candidate_schema_validator_success() -> None:
    # A valid candidate matching candidate_schema.json
    record = {
        "candidate_id": "CAND_1234567",
        "profile": {
            "anonymized_name": "John Doe",
            "headline": "AI Engineer",
            "summary": "Building retrieval systems.",
            "location": "Pune",
            "country": "India",
            "years_of_experience": 5.5,
            "current_title": "Software Engineer",
            "current_company": "Acme Corp",
            "current_company_size": "11-50",
            "current_industry": "Tech",
        },
        "career_history": [
            {
                "company": "Acme Corp",
                "title": "Software Engineer",
                "start_date": "2020-01-01",
                "end_date": None,
                "duration_months": 36,
                "is_current": True,
                "industry": "Tech",
                "company_size": "11-50",
                "description": "Building NLP search systems.",
            }
        ],
        "education": [
            {
                "institution": "IIT",
                "degree": "B.Tech",
                "field_of_study": "Computer Science",
                "start_year": 2016,
                "end_year": 2020,
                "grade": "8.5",
                "tier": "tier_1",
            }
        ],
        "skills": [
            {
                "name": "Python",
                "proficiency": "expert",
                "endorsements": 10,
                "duration_months": 48,
            }
        ],
        "redrob_signals": {
            "profile_completeness_score": 90.0,
            "signup_date": "2020-01-01",
            "last_active_date": "2026-06-20",
            "open_to_work_flag": True,
            "profile_views_received_30d": 12,
            "applications_submitted_30d": 2,
            "recruiter_response_rate": 0.8,
            "avg_response_time_hours": 2.5,
            "skill_assessment_scores": {"Python": 95.0},
            "connection_count": 150,
            "endorsements_received": 25,
            "notice_period_days": 15,
            "expected_salary_range_inr_lpa": {"min": 18.0, "max": 24.0},
            "preferred_work_mode": "hybrid",
            "willing_to_relocate": True,
            "github_activity_score": 85.0,
            "search_appearance_30d": 30,
            "saved_by_recruiters_30d": 5,
            "interview_completion_rate": 0.95,
            "offer_acceptance_rate": 1.0,
            "verified_email": True,
            "verified_phone": True,
            "linkedin_connected": True,
        },
    }

    is_valid, reason = validate_against_schema(record)
    assert is_valid, f"Validation failed: {reason}"


def test_candidate_schema_validator_fails_invalid_id() -> None:
    record = {
        "candidate_id": "CAND_123",  # invalid id pattern
        "profile": {},  # missing required fields
        "career_history": [],
        "education": [],
        "skills": [],
        "redrob_signals": {},
    }
    is_valid, reason = validate_against_schema(record)
    assert not is_valid
    assert "candidate_id" in reason or "profile" in reason


def test_jd_structurer_and_pipeline(tmp_path) -> None:
    # 1. Create a dummy docx job description
    doc_path = tmp_path / "job_description.docx"
    doc = docx.Document()
    doc.add_paragraph("Job Description: Senior AI Engineer — Founding Team")
    doc.add_paragraph("Things you absolutely need:")
    doc.add_paragraph("- Production experience with embeddings-based retrieval systems")
    doc.add_paragraph(
        "- Production experience with vector databases or hybrid search infrastructure"
    )
    doc.add_paragraph("- Strong Python")
    doc.add_paragraph("- Hands-on experience designing evaluation frameworks for ranking systems")
    doc.add_paragraph("Things we'd like you to have:")
    doc.add_paragraph("- LLM fine-tuning experience (LoRA, QLoRA, PEFT)")
    doc.save(doc_path)

    # 2. Create a dummy candidates.jsonl file
    candidates_path = tmp_path / "candidates.jsonl"
    record = {
        "candidate_id": "CAND_1234567",
        "profile": {
            "anonymized_name": "John Doe",
            "headline": "AI Engineer",
            "summary": "Building retrieval systems.",
            "location": "Pune",
            "country": "India",
            "years_of_experience": 5.5,
            "current_title": "Software Engineer",
            "current_company": "Acme Corp",
            "current_company_size": "11-50",
            "current_industry": "Tech",
        },
        "career_history": [
            {
                "company": "Acme Corp",
                "title": "Software Engineer",
                "start_date": "2020-01-01",
                "end_date": None,
                "duration_months": 36,
                "is_current": True,
                "industry": "Tech",
                "company_size": "11-50",
                "description": "Building NLP search systems.",
            }
        ],
        "education": [],
        "skills": [
            {
                "name": "Python",
                "proficiency": "expert",
                "endorsements": 10,
                "duration_months": 48,
            }
        ],
        "redrob_signals": {
            "profile_completeness_score": 90.0,
            "signup_date": "2020-01-01",
            "last_active_date": "2026-06-20",
            "open_to_work_flag": True,
            "profile_views_received_30d": 12,
            "applications_submitted_30d": 2,
            "recruiter_response_rate": 0.8,
            "avg_response_time_hours": 2.5,
            "skill_assessment_scores": {"Python": 95.0},
            "connection_count": 150,
            "endorsements_received": 25,
            "notice_period_days": 15,
            "expected_salary_range_inr_lpa": {"min": 18.0, "max": 24.0},
            "preferred_work_mode": "hybrid",
            "willing_to_relocate": True,
            "github_activity_score": 85.0,
            "search_appearance_30d": 30,
            "saved_by_recruiters_30d": 5,
            "interview_completion_rate": 0.95,
            "offer_acceptance_rate": 1.0,
            "verified_email": True,
            "verified_phone": True,
            "linkedin_connected": True,
        },
    }
    with open(candidates_path, "w", encoding="utf-8") as f:
        f.write(json.dumps(record) + "\n")

    # 3. Run Pipeline
    processed_dir = tmp_path / "processed"
    pipeline = Phase2Pipeline()
    result = pipeline.run(tmp_path, processed_dir)

    assert result.candidate_count == 1
    assert result.job_count == 1
    assert result.rejected_count == 0

    # Verify JSON JobRecord exists
    job_json_file = processed_dir / "job_redrob_senior_ai_engineer.json"
    assert job_json_file.exists()
    with open(job_json_file, encoding="utf-8") as f:
        job_data = json.load(f)
    assert job_data["id"] == "job_redrob_senior_ai_engineer"
    assert "embeddings-based retrieval systems" in job_data["must_have_skills"]

    # Verify flattened candidate parquet contains engineered scalars
    cand_df = pd.read_parquet(result.candidates_path)
    assert len(cand_df) == 1
    assert cand_df.iloc[0]["id"] == "CAND_1234567"
    assert cand_df.iloc[0]["total_career_months"] == 36
    assert cand_df.iloc[0]["num_employers"] == 1
    assert bool(cand_df.iloc[0]["current_role_ai_relevant"])
